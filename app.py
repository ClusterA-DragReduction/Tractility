import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import re
import tempfile
import os
import json
import matplotlib.font_manager as fm
from openpyxl import load_workbook

# ==================== 全局配置 ====================
st.set_page_config(
    page_title="拉伸测试报告生成器",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ==================== 字体设置函数 ====================
def setup_matplotlib_font():
    import platform
    if platform.system() == 'Windows':
        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
    elif platform.system() == 'Linux':
        plt.rcParams['font.sans-serif'] = ['WenQuanYi Micro Hei', 'Noto Sans CJK SC']
    else:
        plt.rcParams['font.sans-serif'] = ['STHeiti', 'Arial Unicode MS']
    plt.rcParams['axes.unicode_minus'] = False


# ==================== 原始数据解析函数（通用） ====================
GROUP_STRIDE = 11
DATA_START_ROW = 10
HEADER_ROW_TESTID = 0
HEADER_ROW_DIMS = 2

COL_FORCE = 1  # 荷重(kgf) 列索引
COL_DISP = 2  # 位移(mm) 列索引
COL_STRESS = 9  # 应力(MPa) 列索引（拉伸测试专用）
COL_STRAIN = 10  # 应变(%) 列索引（拉伸测试专用）

COL_WIDTH = 1
COL_THICK = 3
COL_GAUGE = 1
COL_AREA = 7


def find_all_test_groups(df_raw):
    groups = []
    for col_idx in range(df_raw.shape[1]):
        cell_val = df_raw.iloc[HEADER_ROW_TESTID, col_idx]
        if pd.notna(cell_val) and "测试编号:" in str(cell_val):
            groups.append((col_idx, str(cell_val).strip()))
    return groups


def extract_group_dimensions(df_raw, group_offset):
    area_raw = df_raw.iloc[HEADER_ROW_DIMS, group_offset + COL_AREA]
    area = float(area_raw) if pd.notna(area_raw) else 0.0
    width_raw = df_raw.iloc[HEADER_ROW_DIMS + 1, group_offset + COL_WIDTH]
    thickness_raw = df_raw.iloc[HEADER_ROW_DIMS + 1, group_offset + COL_THICK]
    width = float(width_raw) if pd.notna(width_raw) else 0.0
    thickness = float(thickness_raw) if pd.notna(thickness_raw) else 0.0
    gauge_raw = df_raw.iloc[HEADER_ROW_DIMS + 2, group_offset + COL_GAUGE]
    gauge_length = float(gauge_raw) if pd.notna(gauge_raw) else 50.0
    return width, thickness, gauge_length, area


def extract_group_data(df_raw, group_offset, test_type):
    fc = group_offset + COL_FORCE
    dc = group_offset + COL_DISP
    force_kgf = pd.to_numeric(df_raw.iloc[DATA_START_ROW:, fc], errors='coerce').dropna().values
    disp = pd.to_numeric(df_raw.iloc[DATA_START_ROW:, dc], errors='coerce').dropna().values
    if test_type == "拉伸性能测试":
        sc = group_offset + COL_STRESS
        stc = group_offset + COL_STRAIN
        stress = pd.to_numeric(df_raw.iloc[DATA_START_ROW:, sc], errors='coerce').dropna().values
        strain = pd.to_numeric(df_raw.iloc[DATA_START_ROW:, stc], errors='coerce').dropna().values
        min_len = min(len(force_kgf), len(disp), len(stress), len(strain))
        return force_kgf[:min_len], disp[:min_len], stress[:min_len], strain[:min_len]
    else:
        min_len = min(len(force_kgf), len(disp))
        return force_kgf[:min_len], disp[:min_len], None, None


# ==================== 性能计算函数 ====================
def calculate_mechanical_properties(force_kgf, disp, stress, strain, gauge_length, area):
    if len(force_kgf) == 0:
        return {k: 0.0 for k in ["max_force_N", "max_disp", "max_strain_pct",
                                 "tensile_strength", "E_modulus", "yield_stress", "yield_strain",
                                 "break_stress", "break_strain"]} | {"gauge_length": gauge_length, "area": area}
    max_force_N = np.max(force_kgf) * 9.80665
    idx_max = np.argmax(force_kgf)
    max_disp = disp[idx_max]
    max_strain_pct = max_disp / gauge_length * 100.0 if gauge_length > 0 else 0.0
    tensile_strength = np.max(stress)
    si = np.argsort(strain)
    ss = stress[si]
    sp = strain[si]
    sd = sp / 100.0
    break_stress = float(ss[-1])
    break_strain = float(sp[-1])
    mask = sd <= 0.01
    if np.sum(mask) < 3:
        mask = np.arange(len(sd)) < int(0.01 * len(sd))
    x = sd[mask]
    y = ss[mask]
    E_modulus = np.polyfit(x, y, 1)[0] if len(x) > 1 else 0.0
    if E_modulus > 0 and len(ss) > 1:
        off = 0.002
        off_line = E_modulus * (sd - off)
        diff = ss - off_line
        sc = np.where(np.diff(np.sign(diff)) != 0)[0]
        if len(sc) > 0:
            idx = sc[0]
            x1, x2 = sd[idx], sd[idx + 1]
            y1, y2 = diff[idx], diff[idx + 1]
            if y2 != y1:
                t = -y1 / (y2 - y1)
                ys_d = x1 + t * (x2 - x1)
                fi = interp1d(sd, ss, kind='linear', fill_value='extrapolate')
                yield_stress = float(fi(ys_d))
            else:
                yield_stress = float(ss[idx])
                ys_d = float(sd[idx])
        else:
            yield_stress = tensile_strength
            ys_d = float(sd[-1])
    else:
        yield_stress = tensile_strength
        ys_d = float(sd[-1]) if len(sd) > 0 else 0.0
    return {
        "max_force_N": max_force_N, "max_disp": max_disp,
        "max_strain_pct": max_strain_pct,
        "tensile_strength": tensile_strength, "E_modulus": E_modulus,
        "yield_stress": yield_stress, "yield_strain": ys_d * 100.0,
        "break_stress": break_stress, "break_strain": break_strain,
        "gauge_length": gauge_length, "area": area,
    }


def calculate_peel_properties(force_kgf, disp, width_mm):
    if len(force_kgf) == 0:
        return {"peel_strength_gf_cm": 0, "avg_force_N": 0, "max_force_N": 0, "min_force_N": 0}
    n = len(force_kgf)
    start = int(n * 0.2)
    end = int(n * 0.8)
    if end <= start:
        start = 0
        end = n
    stable_force_kgf = force_kgf[start:end]
    avg_force_kgf = np.mean(stable_force_kgf)
    max_force_kgf = np.max(force_kgf)
    min_force_kgf = np.min(force_kgf)
    width_cm = width_mm / 10.0
    peel_strength_gf_cm = avg_force_kgf * 1000 / width_cm if width_cm > 0 else 0
    return {
        "peel_strength_gf_cm": peel_strength_gf_cm,
        "avg_force_N": avg_force_kgf * 9.80665,
        "max_force_N": max_force_kgf * 9.80665,
        "min_force_N": min_force_kgf * 9.80665
    }


def parse_test_id(full_text):
    text = full_text
    if "测试编号:" in text:
        text = text.split("测试编号:", 1)[-1].strip()
    ts_match = re.search(r'\s+(\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{2}:\d{2}(?:\.\d+)?)\s*$', text)
    timestamp = ""
    clean_id = text
    if ts_match:
        timestamp = ts_match.group(1).strip()
        clean_id = text[:ts_match.start()].strip()
    batch_no = re.sub(r'-\d+\s*$', '', clean_id)
    return clean_id, batch_no, timestamp


# ==================== 编辑数据相关函数 ====================
def get_current_data_for_group(test_id):
    if test_id in st.session_state.edited_data:
        df = st.session_state.edited_data[test_id]
        if st.session_state.test_type == "拉伸性能测试":
            return {
                "force": df["力(kgf)"].values,
                "disp": df["位移(mm)"].values,
                "stress": df["应力(MPa)"].values,
                "strain": df["应变(%)"].values,
            }
        else:
            return {
                "force": df["力(kgf)"].values,
                "disp": df["位移(mm)"].values,
                "stress": None,
                "strain": None,
            }
    else:
        raw = st.session_state.raw_group_data[test_id]
        return raw


def recalc_all_properties():
    new_props = []
    for test_id in st.session_state.test_ids:
        data = get_current_data_for_group(test_id)
        gauge = st.session_state.raw_group_data[test_id]["gauge_length"]
        area = st.session_state.raw_group_data[test_id]["area"]
        width = st.session_state.raw_group_data[test_id]["width"]
        if st.session_state.test_type == "拉伸性能测试":
            props = calculate_mechanical_properties(
                data["force"], data["disp"], data["stress"], data["strain"],
                gauge, area
            )
        else:
            props = calculate_peel_properties(data["force"], data["disp"], width)
        new_props.append(props)
    st.session_state.all_props = new_props
    return new_props


def reset_data_for_group(test_id):
    raw = st.session_state.raw_group_data[test_id]
    if st.session_state.test_type == "拉伸性能测试":
        df = pd.DataFrame({
            "力(kgf)": raw["force"],
            "位移(mm)": raw["disp"],
            "应力(MPa)": raw["stress"] if raw["stress"] is not None else [],
            "应变(%)": raw["strain"] if raw["strain"] is not None else []
        })
    else:
        df = pd.DataFrame({
            "力(kgf)": raw["force"],
            "位移(mm)": raw["disp"]
        })
    st.session_state.edited_data[test_id] = df
    recalc_all_properties()


# ==================== 导出/导入功能 ====================
def export_edited_data():
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        meta_data = []
        for test_id in st.session_state.test_ids:
            gauge = st.session_state.raw_group_data[test_id]["gauge_length"]
            area = st.session_state.raw_group_data[test_id]["area"]
            width = st.session_state.raw_group_data[test_id].get("width", 0)
            meta_data.append({"测试组ID": test_id, "标距(mm)": gauge, "面积(mm²)": area, "宽度(mm)": width})
        df_meta = pd.DataFrame(meta_data)
        df_meta.to_excel(writer, sheet_name="元数据", index=False)
        for test_id in st.session_state.test_ids:
            df = st.session_state.edited_data[test_id]
            df.to_excel(writer, sheet_name=test_id, index=False)
    output.seek(0)
    return output


def import_edited_data(uploaded_file):
    try:
        xlsx = pd.ExcelFile(uploaded_file)
        if "元数据" not in xlsx.sheet_names:
            st.error("无效文件：缺少“元数据”工作表")
            return False
        df_meta = pd.read_excel(uploaded_file, sheet_name="元数据")
        required_cols = ["测试组ID", "标距(mm)", "面积(mm²)"]
        if not all(col in df_meta.columns for col in required_cols):
            st.error("元数据工作表缺少必要列")
            return False

        test_ids = []
        raw_group_data = {}
        edited_data = {}

        for _, row in df_meta.iterrows():
            tid = row["测试组ID"]
            gauge = row["标距(mm)"]
            area = row["面积(mm²)"]
            width = row.get("宽度(mm)", 0)
            if tid not in xlsx.sheet_names:
                st.warning(f"测试组 {tid} 没有数据工作表，跳过")
                continue
            df_data = pd.read_excel(uploaded_file, sheet_name=tid)
            if st.session_state.test_type == "拉伸性能测试":
                required_data_cols = ["力(kgf)", "位移(mm)", "应力(MPa)", "应变(%)"]
                if not all(col in df_data.columns for col in required_data_cols):
                    st.warning(f"测试组 {tid} 的数据工作表列名不正确，跳过")
                    continue
                raw_group_data[tid] = {
                    "force": df_data["力(kgf)"].values,
                    "disp": df_data["位移(mm)"].values,
                    "stress": df_data["应力(MPa)"].values,
                    "strain": df_data["应变(%)"].values,
                    "gauge_length": gauge,
                    "area": area,
                    "width": width
                }
            else:
                required_data_cols = ["力(kgf)", "位移(mm)"]
                if not all(col in df_data.columns for col in required_data_cols):
                    st.warning(f"测试组 {tid} 的数据工作表列名不正确，跳过")
                    continue
                raw_group_data[tid] = {
                    "force": df_data["力(kgf)"].values,
                    "disp": df_data["位移(mm)"].values,
                    "stress": None,
                    "strain": None,
                    "gauge_length": gauge,
                    "area": area,
                    "width": width
                }
            edited_data[tid] = df_data.copy()
            test_ids.append(tid)

        if len(test_ids) == 0:
            st.error("未找到任何有效的测试组数据")
            return False

        st.session_state.test_ids = test_ids
        st.session_state.raw_group_data = raw_group_data
        st.session_state.edited_data = edited_data
        st.session_state.all_props = recalc_all_properties()
        first_id = test_ids[0]
        batch_no_match = re.sub(r'-\d+$', '', first_id)
        st.session_state.batch_no = batch_no_match
        st.session_state.timestamp = "导入数据"
        if "temp_edited_df" in st.session_state:
            del st.session_state.temp_edited_df
        return True
    except Exception as e:
        st.error(f"导入失败: {str(e)}")
        return False


# ==================== 绘图函数 ====================
def plot_custom_chart(selected_groups, x_var_key, y_var_key, filter_method,
                      strain_min, strain_max, row_start, row_end,
                      x_label, y_label, line_width, color_mode, custom_colors):
    setup_matplotlib_font()
    fig, ax = plt.subplots(figsize=(8, 5))
    var_map_inv = {"位移 (mm)": "disp", "力 (kgf)": "force"}
    if st.session_state.test_type == "拉伸性能测试":
        var_map_inv.update({"应力 (MPa)": "stress", "应变 (%)": "strain"})
    x_field = var_map_inv[x_var_key]
    y_field = var_map_inv[y_var_key]

    for idx, test_id in enumerate(selected_groups):
        data = get_current_data_for_group(test_id)
        x_raw = data[x_field]
        y_raw = data[y_field]
        if x_raw is None or y_raw is None:
            continue

        if filter_method == "按应变范围" and st.session_state.test_type == "拉伸性能测试":
            strain_arr = data["strain"]
            mask = (strain_arr >= strain_min) & (strain_arr <= strain_max)
            x_plot = x_raw[mask]
            y_plot = y_raw[mask]
        elif filter_method == "按行号范围":
            start = max(0, row_start)
            end = min(len(x_raw) - 1, row_end)
            x_plot = x_raw[start:end + 1]
            y_plot = y_raw[start:end + 1]
        else:
            x_plot = x_raw
            y_plot = y_raw

        if len(x_plot) == 0:
            continue

        if color_mode == "自定义每个测试组" and test_id in custom_colors:
            color = custom_colors[test_id]
        else:
            color = plt.cm.tab10(idx % 10)

        ax.plot(x_plot, y_plot, color=color, linewidth=line_width, label=test_id)

    ax.set_xlabel(x_label, fontsize=12)
    ax.set_ylabel(y_label, fontsize=12)
    ax.set_title(f"{y_label} - {x_label} 曲线")
    ax.legend(fontsize=9)
    ax.grid(True, alpha=0.3)
    if len(ax.lines) > 0:
        all_y = np.concatenate([line.get_ydata() for line in ax.lines])
        if np.all(all_y >= 0):
            ax.set_ylim(bottom=0)
        all_x = np.concatenate([line.get_xdata() for line in ax.lines])
        if np.all(all_x >= 0):
            ax.set_xlim(left=0)
    return fig


# ==================== Word 报告生成（适配两种测试类型） ====================
def generate_word_report_bytes(selected_groups, x_var, y_var, filter_method,
                               strain_min, strain_max, row_start, row_end,
                               x_label, y_label, line_width, color_mode, custom_colors):
    """生成Word报告并返回BytesIO对象"""
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        chart_path = tmp.name
        fig = plot_custom_chart(
            selected_groups, x_var, y_var, filter_method,
            strain_min, strain_max, row_start, row_end,
            x_label, y_label, line_width, color_mode, custom_colors
        )
        fig.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close(fig)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("XXXX有限公司")
    run.bold = True
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("XXXX测试报告")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()

    header_table = doc.add_table(rows=5, cols=2)
    header_table.style = 'Table Grid'
    header_data = [("测试批号:", st.session_state.batch_no), ("测试人员:", ""), ("客户名称:", ""),
                   ("测试标准:", ""), ("测试日期:", st.session_state.timestamp)]
    for i, (label, val) in enumerate(header_data):
        for j, txt in enumerate([label, val]):
            c = header_table.rows[i].cells[j]
            c.text = ""
            rr = c.paragraphs[0].add_run(txt)
            rr.font.size = Pt(11)
            if j == 0:
                rr.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("测试结果：")
    run.bold = True

    if st.session_state.test_type == "拉伸性能测试":
        headers = ["测试编号", "最大荷重(N)", "最大荷重位移(mm)", "最大荷重伸长率(%)",
                   "抗拉强度(MPa)", "弹性模量(Ei)(MPa)", "屈服强度(MPa)", "屈服伸长率(%)",
                   "断裂强度(MPa)", "断裂伸长率(%)", "标距(mm)", "面积(mm²)"]
        num_data_rows = len(st.session_state.all_props)
        total_rows = 1 + num_data_rows + 3
        table = doc.add_table(rows=total_rows, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for idx, (props, test_id) in enumerate(zip(st.session_state.all_props, st.session_state.test_ids)):
            row = table.rows[1 + idx]
            values = [
                test_id,
                f"{props['max_force_N']:.3f}", f"{props['max_disp']:.3f}",
                f"{props['max_strain_pct']:.3f}", f"{props['tensile_strength']:.3f}",
                f"{props['E_modulus']:.3f}", f"{props['yield_stress']:.3f}",
                f"{props['yield_strain']:.3f}", f"{props['break_stress']:.3f}",
                f"{props['break_strain']:.3f}", f"{props['gauge_length']:.3f}",
                f"{props['area']:.3f}",
            ]
            for j, val in enumerate(values):
                row.cells[j].text = val
                for paragraph in row.cells[j].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stat_rows = [("最大值 Max", np.max), ("最小值 Min", np.min), ("平均值 X-bar", np.mean)]
        stat_fields = ["max_force_N", "max_disp", "max_strain_pct",
                       "tensile_strength", "E_modulus", "yield_stress",
                       "yield_strain", "break_stress", "break_strain"]
        for s_idx, (label, func) in enumerate(stat_rows):
            row = table.rows[1 + num_data_rows + s_idx]
            row.cells[0].text = label
            for f_idx, field in enumerate(stat_fields):
                vals = [p[field] for p in st.session_state.all_props]
                stat_val = func(vals)
                row.cells[1 + f_idx].text = f"{stat_val:.3f}"
            gauge_vals = [p["gauge_length"] for p in st.session_state.all_props]
            row.cells[len(headers) - 2].text = f"{func(gauge_vals):.3f}"
            area_vals = [p["area"] for p in st.session_state.all_props]
            row.cells[len(headers) - 1].text = f"{func(area_vals):.3f}"
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
    else:  # 背胶耐候性测试
        headers = ["测试编号", "180°剥离平均强度(gf/cm)", "宽度(cm)", "单区间最大荷重(N)",
                   "单区间最小荷重(N)", "单区间荷重平均值(N)"]
        num_data_rows = len(st.session_state.all_props)
        total_rows = 1 + num_data_rows + 3  # 加三个统计行
        table = doc.add_table(rows=total_rows, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        # 数据行
        for idx, (props, test_id) in enumerate(zip(st.session_state.all_props, st.session_state.test_ids)):
            width_cm = st.session_state.raw_group_data[test_id]["width"] / 10.0
            row = table.rows[1 + idx]
            values = [
                test_id,
                f"{props['peel_strength_gf_cm']:.3f}",
                f"{width_cm:.3f}",
                f"{props['max_force_N']:.3f}",
                f"{props['min_force_N']:.3f}",
                f"{props['avg_force_N']:.3f}"
            ]
            for j, val in enumerate(values):
                row.cells[j].text = val
                for paragraph in row.cells[j].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 统计行
        stat_rows = [("最大值 Max", np.max), ("最小值 Min", np.min), ("平均值 X-bar", np.mean)]
        stat_fields = ["peel_strength_gf_cm", "max_force_N", "min_force_N", "avg_force_N"]
        for s_idx, (label, func) in enumerate(stat_rows):
            row = table.rows[1 + num_data_rows + s_idx]
            row.cells[0].text = label
            for f_idx, field in enumerate(stat_fields):
                vals = [p[field] for p in st.session_state.all_props]
                stat_val = func(vals)
                row.cells[1 + f_idx].text = f"{stat_val:.3f}"
            # 宽度列单独处理（取平均）
            width_vals = [st.session_state.raw_group_data[t]["width"] / 10.0 for t in st.session_state.test_ids]
            row.cells[len(headers) - 2].text = f"{func(width_vals):.3f}"
            # 最后一列留给平均值，已经填充了avg_force_N，不重复
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("测试曲线：")
    run.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"图表说明：X轴为{x_label}，Y轴为{y_label}，数据筛选方式：{filter_method}")

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    os.unlink(chart_path)
    return word_buffer


# ==================== HTML 报告生成（适配两种测试类型） ====================
def generate_html_report(test_ids, all_props, batch_no, timestamp,
                         x_var_key, y_var_key, x_label, y_label, filter_method,
                         strain_min, strain_max, row_start, row_end,
                         selected_groups, line_width, custom_colors):
    series_data = []
    var_map_inv = {"位移 (mm)": "disp", "力 (kgf)": "force"}
    if st.session_state.test_type == "拉伸性能测试":
        var_map_inv.update({"应力 (MPa)": "stress", "应变 (%)": "strain"})
    x_field = var_map_inv[x_var_key]
    y_field = var_map_inv[y_var_key]

    raw_series_data = []
    for test_id in test_ids:
        data = get_current_data_for_group(test_id)
        if data[x_field] is None or data[y_field] is None:
            continue
        raw_series_data.append({
            "name": test_id,
            "x": data[x_field].tolist(),
            "y": data[y_field].tolist(),
            "strain": data["strain"].tolist() if data["strain"] is not None else []
        })

    if st.session_state.test_type == "拉伸性能测试":
        headers = ["测试编号", "最大荷重(N)", "最大荷重位移(mm)", "最大荷重伸长率(%)",
                   "抗拉强度(MPa)", "弹性模量(Ei)(MPa)", "屈服强度(MPa)", "屈服伸长率(%)",
                   "断裂强度(MPa)", "断裂伸长率(%)", "标距(mm)", "面积(mm²)"]
        table_rows = []
        for idx, test_id in enumerate(test_ids):
            props = all_props[idx]
            row = [
                test_id,
                f"{props['max_force_N']:.3f}", f"{props['max_disp']:.3f}",
                f"{props['max_strain_pct']:.3f}", f"{props['tensile_strength']:.3f}",
                f"{props['E_modulus']:.3f}", f"{props['yield_stress']:.3f}",
                f"{props['yield_strain']:.3f}", f"{props['break_stress']:.3f}",
                f"{props['break_strain']:.3f}", f"{props['gauge_length']:.3f}",
                f"{props['area']:.3f}"
            ]
            table_rows.append(row)
        stat_fields = ["max_force_N", "max_disp", "max_strain_pct",
                       "tensile_strength", "E_modulus", "yield_stress",
                       "yield_strain", "break_stress", "break_strain"]
        stat_rows = []
        for label, func in [("最大值", np.max), ("最小值", np.min), ("平均值", np.mean)]:
            row = [label]
            for field in stat_fields:
                vals = [p[field] for p in all_props]
                row.append(f"{func(vals):.3f}")
            gauge_vals = [p["gauge_length"] for p in all_props]
            row.append(f"{func(gauge_vals):.3f}")
            area_vals = [p["area"] for p in all_props]
            row.append(f"{func(area_vals):.3f}")
            stat_rows.append(row)
    else:
        headers = ["测试编号", "180°剥离平均强度(gf/cm)", "宽度(cm)", "单区间最大荷重(N)",
                   "单区间最小荷重(N)", "单区间荷重平均值(N)"]
        table_rows = []
        for idx, test_id in enumerate(test_ids):
            props = all_props[idx]
            width_cm = st.session_state.raw_group_data[test_id]["width"] / 10.0
            row = [
                test_id,
                f"{props['peel_strength_gf_cm']:.3f}",
                f"{width_cm:.3f}",
                f"{props['max_force_N']:.3f}",
                f"{props['min_force_N']:.3f}",
                f"{props['avg_force_N']:.3f}"
            ]
            table_rows.append(row)
        stat_fields = ["peel_strength_gf_cm", "max_force_N", "min_force_N", "avg_force_N"]
        stat_rows = []
        for label, func in [("最大值", np.max), ("最小值", np.min), ("平均值", np.mean)]:
            row = [label]
            for field in stat_fields:
                vals = [p[field] for p in all_props]
                row.append(f"{func(vals):.3f}")
            width_vals = [st.session_state.raw_group_data[t]["width"] / 10.0 for t in test_ids]
            row.append(f"{func(width_vals):.3f}")
            stat_rows.append(row)

    default_colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22',
                      '#17becf']
    color_map = {}
    for i, tid in enumerate(test_ids):
        if tid in custom_colors:
            color_map[tid] = custom_colors[tid]
        else:
            color_map[tid] = default_colors[i % len(default_colors)]

    html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>测试报告 - {batch_no}</title>
    <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 20px; background: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: auto; background: white; padding: 20px; border-radius: 8px; box-shadow: 0 0 10px rgba(0,0,0,0.1); }}
        h1, h2 {{ color: #2c3e50; text-align: center; }}
        .info {{ background: #e8f4f8; padding: 10px; border-radius: 5px; margin-bottom: 20px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: center; }}
        th {{ background-color: #4CAF50; color: white; }}
        .controls {{ background: #f9f9f9; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
        .control-group {{ display: inline-block; margin-right: 20px; margin-bottom: 10px; }}
        label {{ font-weight: bold; margin-right: 5px; }}
        select, input {{ padding: 5px; border-radius: 3px; border: 1px solid #ccc; }}
        button {{ background: #3498db; color: white; border: none; padding: 8px 16px; border-radius: 4px; cursor: pointer; }}
        button:hover {{ background: #2980b9; }}
        .color-input {{ width: 60px; }}
        .note {{ font-size: 0.9em; color: #555; margin-top: 20px; }}
        hr {{ margin: 15px 0; }}
    </style>
</head>
<body>
<div class="container">
    <h1>XXXX有限公司</h1>
    <h2>XXXX测试报告</h2>
    <div class="info">
        <strong>测试批号:</strong> {batch_no}<br>
        <strong>测试日期:</strong> {timestamp}<br>
        <strong>测试组数:</strong> {len(test_ids)}
    </div>
    <h2>测试结果</h2>
    <table>
        <thead><tr>{''.join(f'<th>{h}</th>' for h in headers)}</thead>
        <tbody>{''.join('<tr>' + ''.join(f'<td>{v}</td>' for v in row) + '</tr>' for row in table_rows)}
        {''.join('<tr>' + ''.join(f'<td><b>{v}</b></td>' for v in row) + '</tr>' for row in stat_rows)}</tbody>
    </table>
    <h2>可交互曲线图</h2>
    <div class="controls">
        <div class="control-group"><label>X轴变量:</label><select id="xVar">
            <option value="disp" {'selected' if x_var_key == '位移 (mm)' else ''}>位移 (mm)</option>
            <option value="force" {'selected' if x_var_key == '力 (kgf)' else ''}>力 (kgf)</option>
            <option value="stress" {'selected' if x_var_key == '应力 (MPa)' and st.session_state.test_type == "拉伸性能测试" else ''}>应力 (MPa)</option>
            <option value="strain" {'selected' if x_var_key == '应变 (%)' and st.session_state.test_type == "拉伸性能测试" else ''}>应变 (%)</option>
        </select></div>
        <div class="control-group"><label>Y轴变量:</label><select id="yVar">
            <option value="disp" {'selected' if y_var_key == '位移 (mm)' else ''}>位移 (mm)</option>
            <option value="force" {'selected' if y_var_key == '力 (kgf)' else ''}>力 (kgf)</option>
            <option value="stress" {'selected' if y_var_key == '应力 (MPa)' and st.session_state.test_type == "拉伸性能测试" else ''}>应力 (MPa)</option>
            <option value="strain" {'selected' if y_var_key == '应变 (%)' and st.session_state.test_type == "拉伸性能测试" else ''}>应变 (%)</option>
        </select></div>
        <div class="control-group"><label>X轴标题:</label><input type="text" id="xLabel" value="{x_label}"></div>
        <div class="control-group"><label>Y轴标题:</label><input type="text" id="yLabel" value="{y_label}"></div>
        <div class="control-group"><label>应变下限 (%):</label><input type="number" id="strainMin" value="{strain_min}" step="0.1"></div>
        <div class="control-group"><label>应变上限 (%):</label><input type="number" id="strainMax" value="{strain_max}" step="0.1"></div>
        <div class="control-group"><label>线条粗细:</label><input type="range" id="lineWidth" min="0.5" max="5" step="0.1" value="{line_width}"></div>
        <div class="control-group"><label>测试组:</label><select id="groupSelect" multiple size="3">
            {''.join(f'<option value="{g}" {"selected" if g in selected_groups else ""}>{g}</option>' for g in test_ids)}
        </select><br><small>按住Ctrl多选</small></div>
        <div><button id="updateBtn">更新图表</button></div>
    </div>
    <div id="plotlyChart"></div>
    <div class="controls" style="margin-top:10px">
        <h4>自定义每组颜色</h4>
        {''.join(f'<div class="control-group"><label>{tid}:</label><input type="color" id="color_{tid}" value="{color_map[tid]}" class="color-input"></div>' for tid in test_ids)}
        <div><button id="applyColorsBtn">应用颜色</button></div>
    </div>
    <div class="note">
        提示：图表支持缩放、平移、下载为PNG；可任意切换X/Y变量、筛选应变范围、选择测试组、调整线条粗细和颜色。<br>
        如需删除或修改数据点，请回到Streamlit应用中使用“高级数据编辑”功能，然后重新导出HTML报告。
    </div>
</div>
<script>
    const allRawData = {json.dumps(raw_series_data, ensure_ascii=False)};
    let currentColors = {json.dumps(color_map, ensure_ascii=False)};
    function applyFilterAndPlot() {{
        const xVar = document.getElementById('xVar').value;
        const yVar = document.getElementById('yVar').value;
        const xLabel = document.getElementById('xLabel').value;
        const yLabel = document.getElementById('yLabel').value;
        const strainMin = parseFloat(document.getElementById('strainMin').value);
        const strainMax = parseFloat(document.getElementById('strainMax').value);
        const selected = Array.from(document.getElementById('groupSelect').selectedOptions).map(opt => opt.value);
        const lineWidth = parseFloat(document.getElementById('lineWidth').value);
        const traces = [];
        for (let testId of selected) {{
            const group = allRawData.find(g => g.name === testId);
            if (!group) continue;
            let xPlot = group.x;
            let yPlot = group.y;
            if (!isNaN(strainMin) && !isNaN(strainMax) && group.strain && group.strain.length) {{
                const mask = group.strain.map(s => s >= strainMin && s <= strainMax);
                xPlot = group.x.filter((_, i) => mask[i]);
                yPlot = group.y.filter((_, i) => mask[i]);
            }}
            const color = currentColors[testId] || '#1f77b4';
            traces.push({{
                x: xPlot,
                y: yPlot,
                mode: 'lines',
                name: testId,
                line: {{ width: lineWidth, color: color }}
            }});
        }}
        const layout = {{
            title: `${{yLabel}} - ${{xLabel}} 曲线`,
            xaxis: {{ title: xLabel }},
            yaxis: {{ title: yLabel }},
            hovermode: 'closest',
            autosize: true,
            margin: {{ l: 50, r: 20, t: 60, b: 50 }}
        }};
        Plotly.newPlot('plotlyChart', traces, layout, {{ responsive: true, displayModeBar: true }});
    }}
    function updateColors() {{
        const testIds = {json.dumps(test_ids, ensure_ascii=False)};
        for (let tid of testIds) {{
            const picker = document.getElementById('color_' + tid);
            if (picker) {{
                currentColors[tid] = picker.value;
            }}
        }}
        applyFilterAndPlot();
    }}
    document.getElementById('updateBtn').addEventListener('click', applyFilterAndPlot);
    document.getElementById('applyColorsBtn').addEventListener('click', updateColors);
    document.getElementById('lineWidth').addEventListener('input', function() {{ applyFilterAndPlot(); }});
    window.onload = applyFilterAndPlot;
</script>
</body>
</html>"""
    return html_template


# ==================== CSS 样式 ====================
st.markdown("""
<style>
    .main-header { font-size: 2rem; font-weight: 600; color: #1E3A8A; text-align: center; margin-bottom: 0.5rem; }
    div[data-testid="stMetric"] label { font-size: 1rem !important; }
    div[data-testid="stMetric"] div { font-size: 1.2rem !important; }
    .stButton button { background-color: #1E3A8A; color: white; border-radius: 8px; font-weight: 500; }
    .stButton button:hover { background-color: #3B82F6; }
    .stDataFrame { border-radius: 12px; overflow: auto; }
    ::-webkit-scrollbar { width: 10px; height: 10px; }
    ::-webkit-scrollbar-track { background: #f1f1f1; border-radius: 10px; }
    ::-webkit-scrollbar-thumb { background: #888; border-radius: 10px; }
    ::-webkit-scrollbar-thumb:hover { background: #555; }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header">📈 拉伸测试数据交互式报告生成器</div>', unsafe_allow_html=True)
st.markdown("上传Excel文件，选择测试类型，自由配置图表，一键生成报告")

# ==================== 初始化 session_state ====================
if "edited_data" not in st.session_state:
    st.session_state.edited_data = {}
if "all_props" not in st.session_state:
    st.session_state.all_props = None
if "raw_group_data" not in st.session_state:
    st.session_state.raw_group_data = None
if "test_ids" not in st.session_state:
    st.session_state.test_ids = []
if "batch_no" not in st.session_state:
    st.session_state.batch_no = ""
if "timestamp" not in st.session_state:
    st.session_state.timestamp = ""
if "test_type" not in st.session_state:
    st.session_state.test_type = "拉伸性能测试"

# ==================== 侧边栏配置 ====================
with st.sidebar:
    st.image("https://img.icons8.com/color/96/000000/test-passed.png", width=80)
    st.markdown("## 配置面板")

    # 测试类型选择（改变时会重置文件上传器）
    test_type = st.radio("测试类型", ["拉伸性能测试", "背胶耐候性测试"],
                         index=0 if st.session_state.test_type == "拉伸性能测试" else 1,
                         key="test_type_radio")
    if test_type != st.session_state.test_type:
        st.session_state.test_type = test_type
        # 清空已有数据，避免类型不匹配
        st.session_state.test_ids = []
        st.session_state.raw_group_data = None
        st.session_state.edited_data = {}
        st.session_state.all_props = None
        # 注意：不调用 st.rerun()，因为下面会重新渲染，但文件上传器会因key变化而重置
        # 强制刷新一次以确保界面更新
        st.rerun()

    # 文件上传器使用动态key，切换测试类型后上传器会重置
    uploaded_file = st.file_uploader(
        "📂 上传Excel文件（支持原始测试数据或编辑后导出的文件）",
        type=["xls", "xlsx"],
        key=f"file_uploader_{st.session_state.test_type}"
    )
    st.markdown("---")
    st.markdown("### 图表自定义")
    if st.session_state.test_type == "拉伸性能测试":
        var_map_display = ["应变 (%)", "应力 (MPa)", ]
    else:
        var_map_display = ["位移 (mm)", "力 (kgf)"]
    x_var = st.selectbox("X轴变量", var_map_display, index=0)
    y_var = st.selectbox("Y轴变量", var_map_display, index=1)
    st.markdown("### 数据筛选")
    if st.session_state.test_type == "拉伸性能测试":
        filter_method = st.radio("筛选方式", ["无筛选", "按应变范围", "按行号范围"], index=0)
    else:
        filter_method = st.radio("筛选方式", ["无筛选", "按行号范围"], index=0)
    st.markdown("### 坐标轴标题")
    # 默认使用包含单位的完整名称
    x_label = st.text_input("X轴标题", x_var)
    y_label = st.text_input("Y轴标题", y_var)
    st.markdown("### 线条样式")
    line_width = st.slider("线宽", 0.5, 5.0, 1.5, step=0.1)
    color_mode = st.radio("颜色模式", ["自动分配", "自定义每个测试组"])

# ==================== 主区域：根据数据源处理 ====================
if uploaded_file is None:
    st.info("👈 请从左侧侧边栏上传Excel文件")
    st.stop()

# 检测文件类型并加载数据（仅在 session_state 中尚无数据时解析）
if not st.session_state.test_ids:
    try:
        xlsx = pd.ExcelFile(uploaded_file)
        is_exported = "元数据" in xlsx.sheet_names
    except:
        is_exported = False

    if is_exported:
        if import_edited_data(uploaded_file):
            st.success("成功导入编辑后的数据文件")
        else:
            st.error("导入失败，请检查文件格式")
            st.stop()
    else:
        df_raw = pd.read_excel(uploaded_file, sheet_name=0, header=None)
        groups = find_all_test_groups(df_raw)
        if not groups:
            st.error("❌ 未找到测试编号，请检查文件格式！")
            st.stop()

        test_ids = []
        all_props_initial = []
        raw_group_data = {}
        for offset, full_text in groups:
            clean_id, batch_no, timestamp = parse_test_id(full_text)
            width, thickness, gauge_length, area = extract_group_dimensions(df_raw, offset)
            force_kgf, disp, stress, strain = extract_group_data(df_raw, offset, st.session_state.test_type)
            if len(force_kgf) == 0:
                continue
            if st.session_state.test_type == "拉伸性能测试":
                props = calculate_mechanical_properties(force_kgf, disp, stress, strain, gauge_length, area)
                raw_group_data[clean_id] = {
                    "force": force_kgf, "disp": disp, "stress": stress, "strain": strain,
                    "gauge_length": gauge_length, "area": area, "width": width
                }
            else:
                props = calculate_peel_properties(force_kgf, disp, width)
                raw_group_data[clean_id] = {
                    "force": force_kgf, "disp": disp, "stress": None, "strain": None,
                    "gauge_length": gauge_length, "area": area, "width": width
                }
            test_ids.append(clean_id)
            all_props_initial.append(props)

        if not test_ids:
            st.error("无有效数据组")
            st.stop()

        st.session_state.test_ids = test_ids
        st.session_state.raw_group_data = raw_group_data
        st.session_state.all_props = all_props_initial
        st.session_state.batch_no = batch_no
        st.session_state.timestamp = timestamp
        for tid in test_ids:
            if tid not in st.session_state.edited_data:
                data = raw_group_data[tid]
                if st.session_state.test_type == "拉伸性能测试":
                    df = pd.DataFrame({
                        "力(kgf)": data["force"],
                        "位移(mm)": data["disp"],
                        "应力(MPa)": data["stress"],
                        "应变(%)": data["strain"]
                    })
                else:
                    df = pd.DataFrame({
                        "力(kgf)": data["force"],
                        "位移(mm)": data["disp"]
                    })
                st.session_state.edited_data[tid] = df

# 如果仍然没有数据（如导入后仍为空），则停止
if not st.session_state.test_ids:
    st.stop()

# ==================== 显示概览指标 ====================
col1, col2, col3 = st.columns(3)
col1.metric("测试批号", st.session_state.batch_no)
col2.metric("测试组数", len(st.session_state.test_ids))
col3.metric("测试日期", st.session_state.timestamp)

with st.expander("📋 测试组详细信息", expanded=False):
    if st.session_state.test_type == "拉伸性能测试":
        df_groups = pd.DataFrame({
            "测试编号": st.session_state.test_ids,
            "标距 (mm)": [st.session_state.raw_group_data[t]["gauge_length"] for t in st.session_state.test_ids],
            "面积 (mm²)": [st.session_state.raw_group_data[t]["area"] for t in st.session_state.test_ids],
            "抗拉强度 (MPa)": [f"{p['tensile_strength']:.1f}" for p in st.session_state.all_props],
            "弹性模量 (MPa)": [f"{p['E_modulus']:.0f}" for p in st.session_state.all_props]
        })
    else:
        df_groups = pd.DataFrame({
            "测试编号": st.session_state.test_ids,
            "宽度 (mm)": [st.session_state.raw_group_data[t]["width"] for t in st.session_state.test_ids],
            "剥离强度 (gf/cm)": [f"{p['peel_strength_gf_cm']:.1f}" for p in st.session_state.all_props],
            "平均荷重 (N)": [f"{p['avg_force_N']:.1f}" for p in st.session_state.all_props]
        })
    st.dataframe(df_groups, use_container_width=True)

# ==================== 高级数据编辑（图表上方，带确认按钮） ====================
st.markdown("---")
st.subheader("✏️ 高级数据编辑")
with st.expander("点击展开/收起数据编辑器", expanded=False):
    edit_group = st.selectbox("选择要编辑的测试组", st.session_state.test_ids, key="edit_group_main")
    if edit_group:
        if "temp_edited_df" not in st.session_state:
            st.session_state.temp_edited_df = st.session_state.edited_data[edit_group].copy()
        if st.session_state.get("last_edit_group") != edit_group:
            st.session_state.temp_edited_df = st.session_state.edited_data[edit_group].copy()
            st.session_state.last_edit_group = edit_group

        st.markdown("**直接修改单元格数值，或删除整行（点击行首删除按钮）**")
        edited_df = st.data_editor(
            st.session_state.temp_edited_df,
            use_container_width=True,
            num_rows="dynamic",
            height=400,
            key=f"data_editor_temp_{edit_group}"
        )
        st.session_state.temp_edited_df = edited_df

        st.markdown("### 快速过滤数据")
        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            filter_col = st.selectbox("选择列", ["力(kgf)", "位移(mm)"], key="filter_col")
        with col_f2:
            filter_op = st.selectbox("操作", ["删除大于阈值", "删除小于阈值", "保留范围"], key="filter_op")
        with col_f3:
            if filter_op == "保留范围":
                min_val = st.number_input("最小值", value=0.0, step=0.1, key="range_min")
                max_val = st.number_input("最大值", value=100.0, step=0.1, key="range_max")
            else:
                threshold = st.number_input("阈值", value=100.0, step=0.1, key="threshold")

        if st.button("应用过滤到临时数据"):
            df_new = st.session_state.temp_edited_df.copy()
            if filter_op == "删除大于阈值":
                df_new = df_new[df_new[filter_col] <= threshold]
            elif filter_op == "删除小于阈值":
                df_new = df_new[df_new[filter_col] >= threshold]
            elif filter_op == "保留范围":
                df_new = df_new[(df_new[filter_col] >= min_val) & (df_new[filter_col] <= max_val)]
            st.session_state.temp_edited_df = df_new
            st.success(f"已过滤，剩余 {len(df_new)} 行。")
            st.rerun()

        col_confirm, col_reset = st.columns(2)
        with col_confirm:
            if st.button("✅ 确认修改并更新图表"):
                st.session_state.edited_data[edit_group] = st.session_state.temp_edited_df.copy()
                recalc_all_properties()
                st.success("数据已更新，性能已重新计算。")
                st.rerun()
        with col_reset:
            if st.button("🔄 重置此组数据"):
                reset_data_for_group(edit_group)
                st.session_state.temp_edited_df = st.session_state.edited_data[edit_group].copy()
                st.rerun()

# ==================== 图表定制与预览 ====================
st.markdown("---")
st.subheader("🎨 图表定制与预览")
selected_groups = st.multiselect("选择要绘制的测试组", st.session_state.test_ids, default=st.session_state.test_ids,
                                 help="可多选")
if not selected_groups:
    st.warning("请至少选择一个测试组")
    st.stop()

custom_colors = {}
if color_mode == "自定义每个测试组":
    st.subheader("自定义线条颜色")
    cols = st.columns(min(4, len(selected_groups)))
    for i, test_id in enumerate(selected_groups):
        with cols[i % 4]:
            custom_colors[test_id] = st.color_picker(f"{test_id}", value="#1f77b4")

strain_min, strain_max = 0.0, 100.0
row_start, row_end = 0, 0
if filter_method == "按应变范围" and st.session_state.test_type == "拉伸性能测试":
    all_strain = np.concatenate([get_current_data_for_group(g)["strain"] for g in selected_groups if
                                 get_current_data_for_group(g)["strain"] is not None])
    if len(all_strain) > 0:
        col_left, col_right = st.columns(2)
        with col_left:
            strain_min = st.number_input("应变下限 (%)", value=float(all_strain.min()), format="%.2f")
        with col_right:
            strain_max = st.number_input("应变上限 (%)", value=float(all_strain.max()), format="%.2f")
elif filter_method == "按行号范围":
    min_len = min([len(get_current_data_for_group(g)["disp"]) for g in selected_groups], default=0)
    if min_len > 0:
        col_left, col_right = st.columns(2)
        with col_left:
            row_start = st.number_input("起始行号", min_value=0, max_value=min_len - 1, value=0, step=1)
        with col_right:
            row_end = st.number_input("结束行号", min_value=row_start, max_value=min_len - 1, value=min_len - 1, step=1)

fig = plot_custom_chart(
    selected_groups, x_var, y_var, filter_method,
    strain_min, strain_max, row_start, row_end,
    x_label, y_label, line_width, color_mode, custom_colors
)
st.pyplot(fig)
plt.close(fig)

# ==================== 导出当前数据按钮 ====================
st.markdown("---")
st.subheader("💾 保存/分享当前编辑状态")
st.download_button(
    label="📥 导出当前所有数据为Excel",
    data=export_edited_data(),
    file_name=f"{st.session_state.batch_no}_编辑后数据.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    use_container_width=True
)

# ==================== 报告生成按钮 ====================
st.markdown("---")
st.subheader("📄 生成报告")
col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
with col_btn2:
    st.download_button(
        label="📄 下载Word报告",
        data=generate_word_report_bytes(
            selected_groups, x_var, y_var, filter_method,
            strain_min, strain_max, row_start, row_end,
            x_label, y_label, line_width, color_mode, custom_colors
        ),
        file_name=f"{st.session_state.batch_no}_测试报告.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

    st.download_button(
        label="🌐 下载可交互HTML报告",
        data=generate_html_report(
            st.session_state.test_ids, st.session_state.all_props, st.session_state.batch_no,
            st.session_state.timestamp,
            x_var, y_var, x_label, y_label, filter_method,
            strain_min, strain_max, row_start, row_end, selected_groups,
            line_width, custom_colors
        ),
        file_name=f"{st.session_state.batch_no}_测试报告.html",
        mime="text/html",
        use_container_width=True
    )